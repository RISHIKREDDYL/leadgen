import os
import re
import shutil
import logging
import argparse
import sys

import fitz                          # PyMuPDF for PDF
import pandas as pd                 # for Excel
from docx import Document           # for .docx

import io
try:
    import tesserocr
except ImportError:
    tesserocr = None
from PIL import Image

# OCR configuration
# Use env var if available, else standard Linux path, or let tesserocr decide
TESSDATA_PREFIX = os.environ.get("TESSDATA_PREFIX", "/usr/share/tesseract-ocr/4.00/tessdata")
LANG            = "eng"
ZOOM            = 2.0    # 2× resolution for higher OCR accuracy

# 1. Configure logger
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s | %(name)s | %(levelname)s | %(message)s"
)
logger = logging.getLogger("leadCategorizer")

def parse_args():
    parser = argparse.ArgumentParser(description="Categorize documents based on lead count.")
    parser.add_argument("--domain", required=True, help="The domain name being processed")
    parser.add_argument("--input_folder", required=True, help="Folder containing downloaded files")
    parser.add_argument("--output_folder", help="Folder to save sorted files (default: files/{domain}/sorted)")
    return parser.parse_args()

# 5a. Extraction function with OCR fallback for PDF
def extract_text_from_pdf(path):
    doc = fitz.open(path)
    mat = fitz.Matrix(ZOOM, ZOOM)
    full_text = []

    # If tesserocr is not installed or TESSDATA_PREFIX is invalid, this might fail.
    # We'll try to use it if available.
    ocr_available = False
    if tesserocr:
         # Check if path exists only if we are using a custom prefix.
         # If TESSDATA_PREFIX is default or set by system, PyTessBaseAPI might find it automatically.
         # For robustness, we try-except the API init.
         ocr_available = True

    try:
        api_context = tesserocr.PyTessBaseAPI(path=TESSDATA_PREFIX, lang=LANG) if ocr_available else None
    except Exception as e:
        logger.warning(f"Failed to initialize Tesseract API: {e}. OCR will be disabled.")
        api_context = None

    if api_context:
        with api_context as ocr_api:
            for page_number in range(len(doc)):
                page = doc.load_page(page_number)
                text = page.get_text("text") or ""
                if text.strip():
                    full_text.append(text)
                else:
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img_bytes = pix.tobytes("png")
                    pil_img = Image.open(io.BytesIO(img_bytes)).convert("L")
                    ocr_api.SetImage(pil_img)
                    ocr_text = ocr_api.GetUTF8Text().strip()
                    full_text.append(ocr_text)
    else:
        # Fallback without OCR
        for page_number in range(len(doc)):
            page = doc.load_page(page_number)
            text = page.get_text("text") or ""
            full_text.append(text)

    doc.close()
    return "\n".join(full_text)

# 5b. Extraction for Excel
def extract_text_from_xlsx(path):
    chunks = []
    sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    for df in sheets.values():
        vals = df.fillna("").astype(str).values.flatten()
        chunks.extend(vals.tolist())
    return " ".join(chunks)

# 5c. Extraction for Word (.docx)
def extract_text_from_docx(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)

def main():
    args = parse_args()
    domain = args.domain
    INPUT_FOLDER = args.input_folder

    if args.output_folder:
        OUTPUT_BASE = args.output_folder
    else:
        # Cross-platform default
        OUTPUT_BASE = os.path.join("files", domain, "sorted")

    # 2. Define folders
    SUBFOLDERS = {
        "500": os.path.join(OUTPUT_BASE, "500plus"),
        "300": os.path.join(OUTPUT_BASE, "301to500"),
        "100": os.path.join(OUTPUT_BASE, "101to300"),
        "50":  os.path.join(OUTPUT_BASE, "50to100"),
    }

    # 3. Create output directories
    # Note: INPUT_FOLDER should already exist from the spider, but we can ensure it.
    os.makedirs(INPUT_FOLDER, exist_ok=True)
    for path in SUBFOLDERS.values():
        os.makedirs(path, exist_ok=True)

    # 4. Compile regex for Indian mobile numbers
    mobile_pattern = re.compile(r'(?<!\d)(?:\+91[\-\s]?|0)?[6-9]\d{9}(?!\d)')

    # 6. Process every file under INPUT_FOLDER (including subfolders)
    all_paths = []
    for root, _, files in os.walk(INPUT_FOLDER):
        for name in files:
            all_paths.append(os.path.join(root, name))

    total = len(all_paths)
    logger.info(f"Starting processing for {domain}. Found {total} files in {INPUT_FOLDER}")

    for idx, src in enumerate(all_paths, start=1):
        fname = os.path.basename(src)
        ext   = os.path.splitext(fname)[1].lower()

        # skip any oddball files
        if ext not in (".pdf", ".docx", ".xlsx", ".xls"):
            logger.debug(f"[{idx}/{total}] skipping unsupported: {fname}")
            continue

        try:
            # dispatch to the right extractor
            if ext == ".pdf":
                content = extract_text_from_pdf(src)
            elif ext in (".xlsx", ".xls"):
                content = extract_text_from_xlsx(src)
            else:  # .docx
                content = extract_text_from_docx(src)

            leads = re.findall(mobile_pattern, content)
            count = len(leads)
            msg   = f"[{idx}/{total}] {fname}: {count} leads"

            if count < 50:
                os.remove(src)
                logger.info(msg + " – deleted (<50 leads).")
                continue

            if   count > 500: dest_key = "500"
            elif count > 300: dest_key = "300"
            elif count > 100: dest_key = "100"
            else:              dest_key = "50"

            dest_folder = SUBFOLDERS[dest_key]
            shutil.move(src, os.path.join(dest_folder, fname))
            logger.info(msg + f" – moved to {os.path.basename(dest_folder)}.")

        except Exception as e:
            logger.error(f"[{idx}/{total}] Error processing {fname}: {e}")

if __name__ == "__main__":
    main()
